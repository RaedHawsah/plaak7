import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def create_report():
    doc = Document()
    
    # Font settings
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)

    # Title
    title = doc.add_heading('تقرير مشروع: تطبيق ساموراي للتوصيل', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('\n')
    
    p = doc.add_paragraph('اسم الطالب: ...............................')
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p = doc.add_paragraph('اسم المادة: ...............................')
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p = doc.add_paragraph('تاريخ التسليم: ...............................')
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.add_page_break()

    # Introduction
    h1 = doc.add_heading('1. مقدمة عن المشروع', level=1)
    h1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p1 = doc.add_paragraph(
        'مشروع "ساموراي" هو تطبيق ويب متكامل مخصص لخدمات توصيل الطعام. '
        'تم استلهام فكرة التصميم من السرعة والدقة التي يتميز بها مقاتلو الساموراي، '
        'مما يعكس التزام التطبيق بتوصيل الطلبات للعملاء بأسرع وقت ممكن وبدقة عالية.'
    )
    p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Technologies
    h2 = doc.add_heading('2. التقنيات المستخدمة', level=1)
    h2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p2.add_run('تم بناء المشروع بالاعتماد على تقنيات الويب الأساسية لضمان السرعة والتوافقية بدون استخدام إطارات عمل معقدة، وذلك لإثبات الإلمام الكامل بأساسيات البرمجة:\n').bold = True
    p2.add_run('• HTML5: لبناء الهيكل الأساسي للصفحات (الرئيسية، من نحن، الشروط، سياسة الخصوصية، الخ).\n')
    p2.add_run('• CSS3: لتصميم واجهة المستخدم (UI) وتجربة المستخدم (UX)، واستخدام الانيميشن والتأثيرات الحركية (Transitions, Hover Effects) بالإضافة لدعم الواجهة باللغة العربية (RTL).\n')
    p2.add_run('• JavaScript (Vanilla): للتحكم بالمنطق البرمجي (اللوجيك)، مثل عمل سلة المشتريات، نظام الفلترة والترتيب للمطاعم.\n')
    p2.add_run('• LocalStorage: لمحاكاة عمل قواعد البيانات (Database)، حيث يتم تخزين حسابات المستخدمين وسجل الطلبات محلياً في المتصفح.\n')

    # Features
    h3 = doc.add_heading('3. أبرز مميزات التطبيق', level=1)
    h3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    features = [
        'واجهة مستخدم احترافية تدعم اللغة العربية من اليمين لليسار بالكامل.',
        'نظام فلترة (تصنيفات المطاعم) لتسهيل العثور على الوجبات المطلوبة (شعبي، شاورما، مشاوي، فطور).',
        'نظام ترتيب للوجبات حسب السعر (من الأقل للأعلى والعكس).',
        'سلة مشتريات ديناميكية لحساب الإجمالي وعملية إتمام الطلب (Checkout) متكاملة.',
        'نظام محاكاة لتسجيل الدخول وإنشاء الحسابات، مع حفظ سجل طلبات المستخدم.',
        'أكواد خصم تفاعلية (Promo Codes) يمكن للمستخدم نسخها بنقرة زر واستخدامها.',
        'صفحات فرعية متكاملة (من نحن، الشروط، الدعم الفني، نموذج الانضمام كشريك).'
    ]
    for f in features:
        p_feat = doc.add_paragraph('• ' + f)
        p_feat.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Conclusion
    h4 = doc.add_heading('4. الخاتمة', level=1)
    h4.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p4 = doc.add_paragraph(
        'مشروع "ساموراي" يمثل تطبيقاً عملياً متكاملاً لمفاهيم تطوير تطبيقات الويب الأمامية (Front-End Development) '
        'وكيفية إدارة البيانات محلياً (State Management & LocalStorage). التصميم والتجربة روعي فيهما أن يحاكيا '
        'تطبيقات التوصيل الحقيقية والموجودة في السوق.'
    )
    p4.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.save('Report.docx')
    print("Report.docx generated successfully.")

if __name__ == '__main__':
    create_report()
